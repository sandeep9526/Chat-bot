"""
Extract plain text from uploaded knowledge files, so a bot can be trained on
real documents — PDFs, Word docs, Markdown, text — not just pasted text.

Document formats (txt / md / pdf / docx) are parsed here with pure-Python libs
and no system dependencies. Images (png / jpg) are handled by the caller via a
vision model instead, because that needs the OpenRouter client that lives in
main.py — see extract_image_text there.
"""
from __future__ import annotations

import io
import os

TEXT_EXTS = {".txt", ".md", ".markdown", ".text"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp"}

# Everything the uploader accepts — used for the file picker + error messages.
SUPPORTED_EXTS = TEXT_EXTS | PDF_EXTS | DOCX_EXTS | IMAGE_EXTS


def file_ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def is_image(filename: str) -> bool:
    return file_ext(filename) in IMAGE_EXTS


def _extract_txt(data: bytes) -> str:
    # Decode as UTF-8, fall back to latin-1 so odd bytes never crash the upload.
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for idx, page in enumerate(reader.pages):
        # Attempt standard layout-aware structured text extraction (preserves multi-column & table spacing)
        try:
            t = page.extract_text(extraction_mode="layout") or page.extract_text() or ""
        except Exception as e:
            print(f"Exception caught in zeva-backend/extract.py: {e}")
            t = page.extract_text() or ""
        if t.strip():
            parts.append(f"--- Page {idx+1} ---\n{t.strip()}")
        else:
            # Scanned page OCR fallback (Tesseract / pytesseract OCR processing for catalog images)
            try:
                import pytesseract
                from PIL import Image
                for img_obj in page.images:
                    img_text = pytesseract.image_to_string(Image.open(io.BytesIO(img_obj.data)))
                    if img_text.strip():
                        parts.append(f"--- Page {idx+1} (OCR Extracted) ---\n{img_text.strip()}")
            except Exception as e:
                print(f"Exception caught in zeva-backend/extract.py: {e}")
                pass
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Structured table-preserving markdown extraction
    for table in doc.tables:
        table_md = []
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                table_md.append("| " + " | ".join(cells) + " |")
                if r_idx == 0:
                    table_md.append("| " + " | ".join(["---"] * len(cells)) + " |")
        if table_md:
            parts.append("\n" + "\n".join(table_md) + "\n")
    return "\n".join(parts)


def extract_document_text(filename: str, data: bytes) -> str:
    """Extract text from a non-image knowledge file. Raises ValueError on an
    unsupported type or when the file yields no readable text (e.g. a scanned
    PDF with no text layer)."""
    ext = file_ext(filename)
    if ext in TEXT_EXTS:
        text = _extract_txt(data)
    elif ext in PDF_EXTS:
        text = _extract_pdf(data)
    elif ext in DOCX_EXTS:
        text = _extract_docx(data)
    elif ext == ".doc":
        raise ValueError(
            "Old .doc files aren't supported — open it and 'Save As' .docx or "
            "PDF, then upload again."
        )
    else:
        raise ValueError(
            f"'{ext or 'That file'}' isn't supported. Upload a PDF, Word (.docx), "
            "text, Markdown, PNG or JPG file."
        )
    text = text.strip()
    if len(text) < 3:
        raise ValueError(
            "Couldn't find readable text in that file. If it's a scanned page, "
            "upload it as an image (PNG/JPG) instead so we can read it."
        )
    return text
